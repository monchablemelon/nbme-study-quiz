import hashlib
import random
import re
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from supabase import create_client


# ============================================================
# STREAMLIT CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="NBME Study Quiz",
    page_icon="🧬",
    layout="wide",
)


# ============================================================
# FILE LOCATIONS
# ============================================================

# Every .docx file placed in the same GitHub folder as this
# Python file becomes its own selectable "question bank" in
# the sidebar. Files that Word creates temporarily while a
# document is open (e.g. "~$My Doc.docx") are ignored.

APP_DIRECTORY = Path(__file__).parent


def discover_docx_files():

    files = [
        path
        for path in APP_DIRECTORY.glob("*.docx")
        if not path.name.startswith("~$")
    ]

    return sorted(files, key=lambda path: path.name.lower())


# ============================================================
# SUPABASE CONNECTION
# ============================================================

@st.cache_resource
def get_supabase_client():
    """
    Create one Supabase client for the Streamlit server.

    IMPORTANT:
    SUPABASE_SECRET_KEY must be stored in Streamlit Secrets.
    Never put the secret key directly in this Python file.
    """
    return create_client(
        st.secrets["SUPABASE_URL"],
        st.secrets["SUPABASE_SECRET_KEY"],
    )


supabase = get_supabase_client()


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text):
    """
    Clean up text extracted from Word.

    Removes:
    - non-breaking spaces
    - zero-width characters
    - markdown-style formatting
    - excessive whitespace
    """

    if text is None:
        return ""

    text = text.replace("\u00a0", " ")
    text = text.replace("\u200b", "")
    text = text.replace("\u200c", "")
    text = text.replace("\u200d", "")
    text = text.replace("\ufeff", "")

    text = text.replace("**", "")
    text = text.replace("__", "")

    text = text.replace("\t", " ")

    return text.strip()


# ============================================================
# READ WORD DOCUMENT
# ============================================================

def read_word_as_text(file_path):
    """
    Read the Word document and return its text.

    Reads:
    - normal paragraphs
    - text inside tables
    """

    document = Document(file_path)

    chunks = []

    for paragraph in document.paragraphs:
        text = paragraph.text

        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:

            row_parts = []

            for cell in row.cells:
                cell_text = cell.text

                if cell_text:
                    row_parts.append(cell_text)

            if row_parts:
                chunks.append(
                    " | ".join(row_parts)
                )

    full_text = "\n".join(chunks)

    full_text = full_text.replace("\r\n", "\n")
    full_text = full_text.replace("\r", "\n")

    return full_text


# ============================================================
# NORMALIZE DOCUMENT
# ============================================================

def normalize_document_text(text):

    lines = []

    for raw_line in text.split("\n"):

        line = clean_text(raw_line)

        if line:
            lines.append(line)

    return "\n".join(lines)


# ============================================================
# IDENTIFY QUESTION NUMBER
# ============================================================

QUESTION_NUMBER_PATTERN = re.compile(r"(\d+)\s*[\.\)]?")


def parse_question_number(line):

    line = clean_text(line)

    line = line.replace("*", "")

    match = re.fullmatch(
        QUESTION_NUMBER_PATTERN,
        line
    )

    if match:
        return int(match.group(1))

    return None


# ============================================================
# IDENTIFY ANSWER CHOICE
# ============================================================

ANSWER_CHOICE_PATTERN = re.compile(
    r"^([A-E])\s*[\.\)]\s*(.+)$",
    flags=re.IGNORECASE
)


def parse_answer_choice(line):

    line = clean_text(line)

    line = line.replace("*", "")

    match = ANSWER_CHOICE_PATTERN.match(line)

    if not match:
        return None

    letter = match.group(1).upper()
    answer = match.group(2).strip()

    return letter, answer


# ============================================================
# FIND SECTION
# ============================================================

def find_section(text, heading):

    match = re.search(
        rf"(?im)^\s*{re.escape(heading)}\s*$",
        text
    )

    if match:
        return match

    return None


def is_heading_paragraph(text, heading):

    return bool(
        re.fullmatch(
            rf"(?i){re.escape(heading)}",
            text.strip()
        )
    )


# ============================================================
# PARSE QUESTIONS
# ============================================================

def parse_questions(question_text):

    lines = question_text.split("\n")

    questions = {}

    current_number = None
    current_question_lines = []
    current_options = {}

    def save_current_question():

        nonlocal current_number
        nonlocal current_question_lines
        nonlocal current_options

        if current_number is None:
            return

        question_text_combined = " ".join(
            current_question_lines
        ).strip()

        if (
            question_text_combined
            and len(current_options) >= 4
        ):
            questions[current_number] = {
                "question": question_text_combined,
                "options": dict(current_options)
            }

    for raw_line in lines:

        line = clean_text(raw_line)

        if not line:
            continue

        question_number = parse_question_number(line)

        if question_number is not None:

            save_current_question()

            current_number = question_number
            current_question_lines = []
            current_options = {}

            continue

        option = parse_answer_choice(line)

        if option is not None and current_number is not None:

            letter, answer = option
            current_options[letter] = answer

            continue

        if current_number is not None:
            current_question_lines.append(line)

    save_current_question()

    return questions


# ============================================================
# FIND EXPLANATION SECTION
# ============================================================

HEADING_NUMBER_PATTERN = re.compile(
    r"(?im)^\s*(\d+)\.\s+(.+?)\s*$"
)


def extract_explanation_section(full_text, answer_sheet_match):

    text_after_answer_sheet = full_text[
        answer_sheet_match.end():
    ]

    for match in HEADING_NUMBER_PATTERN.finditer(
        text_after_answer_sheet
    ):

        start = match.end()

        preview = text_after_answer_sheet[
            start:start + 2000
        ]

        if re.search(
            r"Correct\s+answer\s*:",
            preview,
            flags=re.IGNORECASE
        ):

            return text_after_answer_sheet[match.start():]

    raise ValueError(
        "Could not find the explanation section."
    )


# ============================================================
# PARSE EXPLANATIONS
# ============================================================

def parse_explanations(explanation_text):

    explanations = {}

    matches = list(
        HEADING_NUMBER_PATTERN.finditer(explanation_text)
    )

    for index, match in enumerate(matches):

        number = int(match.group(1))

        if index + 1 < len(matches):
            end = matches[index + 1].start()
        else:
            end = len(explanation_text)

        block = explanation_text[
            match.start():end
        ].strip()

        answer_match = re.search(
            r"Correct\s+answer\s*:\s*([A-E])",
            block,
            flags=re.IGNORECASE
        )

        if answer_match is None:
            continue

        correct_letter = answer_match.group(1).upper()

        explanations[number] = {
            "correct_letter": correct_letter,
            "explanation": block
        }

    return explanations


# ============================================================
# LOAD COMPLETE QUESTION BANK (single .docx file)
# ============================================================

def load_question_bank(file_path):

    if not file_path.exists():

        raise FileNotFoundError(
            "The Word document could not be found:\n\n"
            f"{file_path}"
        )

    raw_text = read_word_as_text(file_path)

    full_text = normalize_document_text(raw_text)

    questions_match = find_section(full_text, "Questions")
    answer_sheet_match = find_section(full_text, "Answer Sheet")

    if questions_match is None:
        raise ValueError(
            "Could not find 'Questions' in the Word document."
        )

    if answer_sheet_match is None:
        raise ValueError(
            "Could not find 'Answer Sheet' in the Word document."
        )

    question_section = full_text[
        questions_match.end():
        answer_sheet_match.start()
    ]

    questions = parse_questions(question_section)

    explanation_section = extract_explanation_section(
        full_text,
        answer_sheet_match
    )

    explanations = parse_explanations(explanation_section)

    final_questions = []

    for number in sorted(questions):

        if number not in explanations:
            continue

        question = questions[number]
        explanation = explanations[number]

        correct_letter = explanation["correct_letter"]

        if correct_letter not in question["options"]:
            continue

        final_questions.append({
            "number": number,
            "question": question["question"],
            "options": question["options"],
            "correct_letter": correct_letter,
            "correct_answer": question["options"][correct_letter],
            "explanation": explanation["explanation"]
        })

    if not final_questions:
        raise ValueError(
            "No complete questions were found "
            "in the Word document."
        )

    return final_questions


# ============================================================
# CACHE A SINGLE QUESTION BANK FILE
# ============================================================

@st.cache_data
def load_cached_question_bank(file_path, modified_time):
    """
    Cache the parsed Word document.

    modified_time ensures Streamlit reloads the document
    when the Word file changes (including when a question
    is removed by an administrator).
    """

    return load_question_bank(Path(file_path))


# ============================================================
# BANK IDENTIFIERS
# ============================================================

def slugify(name):

    slug = re.sub(r"[^a-zA-Z0-9]+", "_", name).strip("_").lower()

    return slug or "bank"


# ============================================================
# "HARDEST QUESTIONS" — A VIRTUAL, CROSS-BANK PRACTICE DECK
# ============================================================

HARDEST_BANK_ID = "hardest_questions"
HARDEST_BANK_LABEL = "🔥 Hardest Questions"
HARDEST_QUESTION_COUNT = 20

# A question needs at least this many recorded attempts
# (across everyone) before its accuracy is trusted enough to
# count toward the hardest-questions ranking. Otherwise a
# single unlucky guess could make a question look "hardest".
HARDEST_MIN_ATTEMPTS = 3


@st.cache_data(ttl=30)
def compute_global_question_stats():
    """
    Aggregate every recorded answer, across every bank, into
    per-question (bank_id, question_number) attempt/correct
    counts. Cached briefly and shared across everyone's
    session, since this reflects global data, not anything
    participant-specific.
    """

    response = (
        supabase
        .table("answer_logs")
        .select("bank_id, question_number, is_correct")
        .execute()
    )

    rows = response.data or []

    stats = {}

    for row in rows:

        bank_id = row.get("bank_id")
        question_number = row.get("question_number")

        if bank_id is None or question_number is None:
            continue

        key = (bank_id, question_number)

        entry = stats.setdefault(key, {"attempts": 0, "correct": 0})

        entry["attempts"] += 1

        if row.get("is_correct"):
            entry["correct"] += 1

    return stats


def build_hardest_questions_bank():
    """
    Build a virtual, read-only deck made up of the questions —
    pooled across every real bank — with the lowest global
    accuracy, so participants can specifically drill their
    weakest spots. Each question is tagged with the real bank
    it actually came from, so answering/flagging/removing it
    here still routes back to that bank's real data.
    """

    try:
        global_stats = compute_global_question_stats()
    except Exception:
        global_stats = {}

    candidates = []

    for bank in VALID_BANKS:

        bank_id = bank["bank_id"]

        for question in bank["questions"]:

            entry = global_stats.get((bank_id, question["number"]))

            if not entry or entry["attempts"] < HARDEST_MIN_ATTEMPTS:
                continue

            accuracy = entry["correct"] / entry["attempts"]

            candidates.append(
                (accuracy, entry["attempts"], bank_id, question)
            )

    # Lowest accuracy first; ties broken by more attempts
    # (more attempts = more confidence it's genuinely hard).
    candidates.sort(key=lambda item: (item[0], -item[1]))

    chosen = candidates[:HARDEST_QUESTION_COUNT]

    hardest_questions = []

    for accuracy, attempts, source_bank_id, question in chosen:

        tagged_question = dict(question)
        tagged_question["source_bank_id"] = source_bank_id
        tagged_question["source_accuracy"] = accuracy
        tagged_question["source_attempts"] = attempts

        hardest_questions.append(tagged_question)

    return hardest_questions


# ============================================================
# LOAD EVERY QUESTION BANK FOUND NEXT TO THE APP
# ============================================================

def load_all_banks():
    """
    Discover every .docx file in the app's folder and parse
    each one into its own independent question bank. A bank
    that fails to parse is kept in the list (so it still shows
    up, disabled, in the sidebar) but is flagged with an error
    instead of crashing the whole app.
    """

    banks = []
    seen_ids = set()

    for file_path in discover_docx_files():

        display_name = file_path.stem

        base_id = slugify(display_name)
        bank_id = base_id
        suffix = 2

        while bank_id in seen_ids:
            bank_id = f"{base_id}_{suffix}"
            suffix += 1

        seen_ids.add(bank_id)

        try:
            modified_time = file_path.stat().st_mtime_ns

            questions = load_cached_question_bank(
                str(file_path),
                modified_time
            )

            error = None

        except Exception as load_error:
            questions = []
            error = str(load_error)

        banks.append({
            "bank_id": bank_id,
            "display_name": display_name,
            "path": file_path,
            "questions": questions,
            "error": error,
        })

    return banks


# ============================================================
# PARTICIPANT ID
# ============================================================

def create_participant_id(participant_code, app_salt):
    """
    Convert a participant's private code into a
    deterministic identifier. The raw code is not stored.
    """

    value = str(app_salt) + "|" + participant_code.strip()

    return hashlib.sha256(value.encode("utf-8")).hexdigest()


# ============================================================
# DATABASE FUNCTIONS — PARTICIPANTS
# ============================================================

def save_participant(participant_id, display_name):

    response = (
        supabase
        .table("participants")
        .upsert(
            {
                "participant_id": participant_id,
                "display_name": display_name,
            },
            on_conflict="participant_id"
        )
        .execute()
    )

    return response


def get_participant_cycle(participant_id, bank_id):
    """
    Get the participant's current "lap" through one specific
    question bank. Defaults to 1 for a brand-new participant,
    or for a bank they haven't started yet.

    NOTE: this reads from the "participant_progress" table,
    which needs a (participant_id, bank_id) unique constraint —
    see the setup notes for the SQL to create it.
    """

    response = (
        supabase
        .table("participant_progress")
        .select("current_cycle")
        .eq("participant_id", participant_id)
        .eq("bank_id", bank_id)
        .execute()
    )

    rows = response.data or []

    if not rows:
        return 1

    cycle = rows[0].get("current_cycle")

    return cycle if cycle else 1


def set_participant_cycle(participant_id, bank_id, cycle):

    (
        supabase
        .table("participant_progress")
        .upsert(
            {
                "participant_id": participant_id,
                "bank_id": bank_id,
                "current_cycle": cycle,
            },
            on_conflict="participant_id,bank_id"
        )
        .execute()
    )


# ============================================================
# DATABASE FUNCTIONS — ANSWERS
# ============================================================

def save_answer(participant_id, display_name, bank_id, question, selected_letter, cycle):
    """
    Save one answer attempt permanently in Supabase, tagged
    with which bank and which cycle it belongs to.
    """

    is_correct = selected_letter == question["correct_letter"]

    record = {
        "participant_id": participant_id,
        "display_name": display_name,
        "bank_id": bank_id,
        "question_number": question["number"],
        "selected_letter": selected_letter,
        "correct_letter": question["correct_letter"],
        "is_correct": is_correct,
        "cycle": cycle,
        "answered_at": datetime.now(timezone.utc).isoformat(),
    }

    supabase.table("answer_logs").insert(record).execute()

    return is_correct


def get_completed_question_numbers(participant_id, bank_id, cycle):
    """
    Question numbers this participant has already answered
    CORRECTLY in this bank during the given cycle. These are
    excluded from the shuffled queue so they aren't repeated
    until the whole bank has been completed.
    """

    response = (
        supabase
        .table("answer_logs")
        .select("question_number")
        .eq("participant_id", participant_id)
        .eq("bank_id", bank_id)
        .eq("cycle", cycle)
        .eq("is_correct", True)
        .execute()
    )

    rows = response.data or []

    return {row["question_number"] for row in rows}


# ============================================================
# BUILD A PARTICIPANT'S QUEUE (no-repeat, cross-session)
# ============================================================

def build_participant_queue(participant_id, bank_id):
    """
    Figure out where this participant is in their current
    cycle through one specific question bank, and hand back a
    freshly-shuffled queue of only the questions in that bank
    they have not yet answered correctly this cycle.

    If they've already finished every question in the current
    cycle, the cycle is advanced and a brand-new full, shuffled
    queue is returned.
    """

    bank_questions = BANKS_BY_ID[bank_id]["questions"]

    cycle = get_participant_cycle(participant_id, bank_id)

    completed_numbers = get_completed_question_numbers(
        participant_id,
        bank_id,
        cycle
    )

    valid_numbers = {question["number"] for question in bank_questions}

    # Only count completions for questions that still exist —
    # an administrator may have removed one since.
    completed_numbers = completed_numbers & valid_numbers

    remaining = [
        question
        for question in bank_questions
        if question["number"] not in completed_numbers
    ]

    if not remaining:

        cycle += 1

        set_participant_cycle(participant_id, bank_id, cycle)

        completed_numbers = set()

        remaining = bank_questions.copy()

    random.shuffle(remaining)

    return cycle, completed_numbers, remaining


def reset_participant_progress(participant_id, bank_id):
    """
    Manually abandon the current cycle for one bank and start
    a fresh one, even if it isn't finished yet.
    """

    current_cycle = get_participant_cycle(participant_id, bank_id)

    set_participant_cycle(participant_id, bank_id, current_cycle + 1)


def start_queue_in_session(participant_id, bank_id):

    cycle, completed_numbers, remaining = build_participant_queue(
        participant_id,
        bank_id
    )

    st.session_state.selected_bank = bank_id
    st.session_state.selected_bank_total = len(BANKS_BY_ID[bank_id]["questions"])
    st.session_state.cycle = cycle
    st.session_state.base_completed = len(completed_numbers)
    st.session_state.questions = remaining
    st.session_state.current_index = 0
    st.session_state.current_correct = False
    st.session_state.quiz_complete = False
    st.session_state.last_feedback = ""
    st.session_state.last_feedback_type = ""
    st.session_state.show_explanation = False
    st.session_state.flash_type = None
    st.session_state.wrong_answer_streak = 0
    st.session_state.meltdown_popup = False


def start_hardest_queue_in_session(participant_id):
    """
    Load the current top-N hardest questions (pooled across
    every bank) as a one-off, freshly-shuffled practice deck.
    This deck isn't tracked with its own persistent cycle —
    it's simply recomputed from live global stats every time
    it's opened, so the mix of questions can change over time.
    """

    hardest_questions = build_hardest_questions_bank()

    random.shuffle(hardest_questions)

    st.session_state.selected_bank = HARDEST_BANK_ID
    st.session_state.selected_bank_total = len(hardest_questions)
    st.session_state.cycle = 1
    st.session_state.base_completed = 0
    st.session_state.questions = hardest_questions
    st.session_state.current_index = 0
    st.session_state.current_correct = False
    st.session_state.quiz_complete = False
    st.session_state.last_feedback = ""
    st.session_state.last_feedback_type = ""
    st.session_state.show_explanation = False
    st.session_state.flash_type = None
    st.session_state.wrong_answer_streak = 0
    st.session_state.meltdown_popup = False


# ============================================================
# SIDEBAR PROGRESS SUMMARY (used for every bank's subpanel)
# ============================================================

def get_bank_progress_summary(participant_id, bank_id):

    bank_questions = BANKS_BY_ID[bank_id]["questions"]
    total = len(bank_questions)

    if total == 0:
        return 1, 0, 0

    cycle = get_participant_cycle(participant_id, bank_id)

    completed_numbers = get_completed_question_numbers(
        participant_id,
        bank_id,
        cycle
    )

    valid_numbers = {question["number"] for question in bank_questions}
    completed_numbers = completed_numbers & valid_numbers

    return cycle, len(completed_numbers), total


# ============================================================
# DATABASE FUNCTIONS — STATISTICS
# ============================================================

def get_question_stats(participant_id, bank_id, question_number):
    """
    Get historical statistics for ONE specific question in ONE
    specific bank, both for this user (across all their cycles)
    and globally.
    """

    response = (
        supabase
        .table("answer_logs")
        .select("selected_letter, is_correct")
        .eq("participant_id", participant_id)
        .eq("bank_id", bank_id)
        .eq("question_number", question_number)
        .execute()
    )

    rows = response.data or []

    attempts = len(rows)

    correct = sum(
        1 for row in rows if row.get("is_correct") is True
    )

    answer_counts = {"A": 0, "B": 0, "C": 0, "D": 0, "E": 0}

    for row in rows:

        letter = row.get("selected_letter")

        if letter in answer_counts:
            answer_counts[letter] += 1

    user_accuracy = (correct / attempts * 100) if attempts else 0

    global_response = (
        supabase
        .table("answer_logs")
        .select("*", count="exact", head=True)
        .eq("bank_id", bank_id)
        .eq("question_number", question_number)
        .execute()
    )

    global_attempts = global_response.count or 0

    global_correct_response = (
        supabase
        .table("answer_logs")
        .select("*", count="exact", head=True)
        .eq("bank_id", bank_id)
        .eq("question_number", question_number)
        .eq("is_correct", True)
        .execute()
    )

    global_correct = global_correct_response.count or 0

    global_accuracy = (
        (global_correct / global_attempts * 100)
        if global_attempts
        else 0
    )

    return {
        "attempts": attempts,
        "correct": correct,
        "accuracy": user_accuracy,
        "answers": answer_counts,
        "global_attempts": global_attempts,
        "global_correct": global_correct,
        "global_accuracy": global_accuracy,
    }


# ============================================================
# ADMIN FUNCTIONS
# ============================================================

def get_all_answers():

    response = (
        supabase
        .table("answer_logs")
        .select("*")
        .order("answered_at", desc=True)
        .execute()
    )

    return response.data or []


def save_flag(participant_id, display_name, bank_id, question_number):

    record = {
        "participant_id": participant_id,
        "display_name": display_name,
        "bank_id": bank_id,
        "question_number": question_number,
        "flagged_at": datetime.now(timezone.utc).isoformat(),
    }

    supabase.table("flagged_questions").insert(record).execute()


def get_flagged_questions():

    response = (
        supabase
        .table("flagged_questions")
        .select("*")
        .order("flagged_at", desc=True)
        .execute()
    )

    return response.data or []


def clear_flags_for_question(bank_id, question_number):

    (
        supabase
        .table("flagged_questions")
        .delete()
        .eq("bank_id", bank_id)
        .eq("question_number", question_number)
        .execute()
    )


# ============================================================
# REMOVE A QUESTION FROM A WORD DOCUMENT
# ============================================================

def _delete_paragraph(paragraph):

    element = paragraph._element
    element.getparent().remove(element)


def strip_question_from_docx(file_path, question_number):
    """
    Permanently delete a question — its prompt, its answer
    choices, its entry in the plain answer-key list, and its
    explanation block — from the Word document on disk.

    NOTE: this only removes content that lives in normal
    paragraphs (which is how this document is structured).
    If a particular question's choices were placed inside a
    Word table instead, they will need to be removed by hand.
    """

    document = Document(file_path)

    paragraphs = document.paragraphs

    questions_heading_index = None
    answer_sheet_heading_index = None

    for index, paragraph in enumerate(paragraphs):

        text = clean_text(paragraph.text)

        if (
            questions_heading_index is None
            and is_heading_paragraph(text, "Questions")
        ):
            questions_heading_index = index
            continue

        if is_heading_paragraph(text, "Answer Sheet"):
            answer_sheet_heading_index = index
            break

    if questions_heading_index is None or answer_sheet_heading_index is None:
        raise ValueError(
            "Could not locate the document's Questions / "
            "Answer Sheet sections."
        )

    to_delete = []

    # --------------------------------------------------
    # The question prompt + its answer choices
    # --------------------------------------------------

    in_target_question = False

    for index in range(
        questions_heading_index + 1,
        answer_sheet_heading_index
    ):

        paragraph = paragraphs[index]
        text = clean_text(paragraph.text)

        if not text:
            continue

        number_match = re.fullmatch(
            QUESTION_NUMBER_PATTERN,
            text.replace("*", "")
        )

        if number_match:

            in_target_question = (
                int(number_match.group(1)) == question_number
            )

            if in_target_question:
                to_delete.append(paragraph)

            continue

        if in_target_question:
            to_delete.append(paragraph)

    # --------------------------------------------------
    # The answer-key entry AND the explanation block
    # (both are "N. ..." headings after "Answer Sheet")
    # --------------------------------------------------

    heading_indices = []

    for index in range(
        answer_sheet_heading_index + 1,
        len(paragraphs)
    ):

        text = clean_text(paragraphs[index].text)

        if not text:
            continue

        if HEADING_NUMBER_PATTERN.fullmatch(text):
            heading_indices.append(index)

    for position, index in enumerate(heading_indices):

        text = clean_text(paragraphs[index].text)
        match = HEADING_NUMBER_PATTERN.fullmatch(text)
        number = int(match.group(1))

        if number != question_number:
            continue

        end_index = (
            heading_indices[position + 1]
            if position + 1 < len(heading_indices)
            else len(paragraphs)
        )

        for inner_index in range(index, end_index):
            to_delete.append(paragraphs[inner_index])

    if not to_delete:
        raise ValueError(
            f"Question {question_number} could not be found "
            "in the Word document."
        )

    for paragraph in to_delete:
        _delete_paragraph(paragraph)

    document.save(file_path)


def remove_question_everywhere(bank_id, question_number):
    """
    Strip a question from its bank's source document and clear
    any reports about it. Raises on failure.
    """

    bank = BANKS_BY_ID[bank_id]

    strip_question_from_docx(bank["path"], question_number)

    try:
        clear_flags_for_question(bank_id, question_number)
    except Exception:
        # Flag cleanup is best-effort; the document edit is
        # what actually matters.
        pass


def drop_question_from_session(question_number):
    """
    Remove a just-deleted question from the participant's
    in-progress queue so it isn't shown again this session.
    Only relevant when the deletion happened in the bank the
    participant currently has active.
    """

    st.session_state.questions = [
        question
        for question in st.session_state.questions
        if question["number"] != question_number
    ]

    st.session_state.show_explanation = False
    st.session_state.current_correct = False
    st.session_state.last_feedback = ""
    st.session_state.last_feedback_type = ""
    st.session_state.flash_type = None

    if st.session_state.current_index >= len(st.session_state.questions):
        st.session_state.quiz_complete = True
        st.session_state.current_index = 0


# ============================================================
# INITIALIZE SESSION STATE
# ============================================================

def initialize_quiz_state():

    defaults = {
        "quiz_started": False,
        "participant_id": None,
        "display_name": "",
        "selected_bank": None,
        "selected_bank_total": 0,
        "questions": [],
        "current_index": 0,
        "current_correct": False,
        "quiz_complete": False,
        "last_feedback": "",
        "last_feedback_type": "",
        "show_explanation": False,
        "flash_type": None,
        "wrong_answer_streak": 0,
        "meltdown_popup": False,

        # Cross-session, no-repeat progress tracking
        # (scoped to whichever bank is currently selected).
        "cycle": 1,
        "base_completed": 0,

        # Flagging / admin. Flags are stored as
        # (bank_id, question_number) tuples since question
        # numbers can repeat across different banks.
        "flagged_this_session": set(),
        "is_admin": False,
    }

    for key, value in defaults.items():

        if key not in st.session_state:
            st.session_state[key] = value


initialize_quiz_state()


# ============================================================
# ANSWER FLASH
# ============================================================

if st.session_state.get("flash_type"):

    if st.session_state.flash_type == "correct":
        flash_color = "rgba(0, 200, 0, 0.40)"
    else:
        flash_color = "rgba(255, 0, 0, 0.40)"

    st.markdown(
        f"""
        <style>

        @keyframes answerFlash {{
            0% {{ opacity: 0; }}
            20% {{ opacity: 1; }}
            60% {{ opacity: 1; }}
            100% {{ opacity: 0; }}
        }}

        .answer-flash {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100vw;
            height: 100vh;
            background: {flash_color};
            z-index: 999999;
            pointer-events: none;
            animation: answerFlash 0.65s ease-out forwards;
        }}

        </style>

        <div class="answer-flash"></div>
        """,
        unsafe_allow_html=True
    )

    st.session_state.flash_type = None


# ============================================================
# MELTDOWN 🫠 POPUP
# ============================================================

if st.session_state.get("meltdown_popup"):

    st.markdown(
        """
        <style>

        @keyframes meltdownFloat {
            0% { opacity: 0; transform: translate(100px, 20px) scale(0.65); }
            15% { opacity: 1; transform: translate(0, 0) scale(1); }
            55% { opacity: 1; transform: translate(-8px, -8px) scale(1.05); }
            75% { opacity: 0.9; transform: translate(-15px, -15px) scale(1.02); }
            100% { opacity: 0; transform: translate(-55px, -45px) scale(0.9); }
        }

        .meltdown-popup {
            position: fixed;
            right: 30px;
            top: 50%;
            z-index: 1000000;
            pointer-events: none;
            font-size: 72px;
            line-height: 1;
            animation: meltdownFloat 2s ease-out forwards;
            filter: drop-shadow(0 5px 12px rgba(0, 0, 0, 0.25));
        }

        </style>

        <div class="meltdown-popup">🫠</div>
        """,
        unsafe_allow_html=True
    )

    st.session_state.meltdown_popup = False


# ============================================================
# LOAD ALL QUESTION BANKS
# ============================================================

ALL_BANKS = load_all_banks()
BANKS_BY_ID = {bank["bank_id"]: bank for bank in ALL_BANKS}
VALID_BANKS = [bank for bank in ALL_BANKS if not bank["error"]]

if not ALL_BANKS:
    st.error(
        "No .docx question bank files were found in the app's "
        "folder. Add one or more Word documents next to this "
        "script in the GitHub repository."
    )
    st.stop()

if not VALID_BANKS:
    st.error("None of the question bank documents could be loaded.")
    for bank in ALL_BANKS:
        st.exception(ValueError(f"{bank['display_name']}: {bank['error']}"))
    st.stop()


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:

    st.title("🧠 NBME Study Quiz")

    if st.session_state.quiz_started:

        st.write(f"**Participant:** {st.session_state.display_name}")

        st.divider()

        st.write("**Practice**")

        hardest_preview = build_hardest_questions_bank()
        hardest_total_preview = len(hardest_preview)

        is_hardest_selected = (
            st.session_state.selected_bank == HARDEST_BANK_ID
        )

        hardest_label = (
            f"{'▶ ' if is_hardest_selected else ''}{HARDEST_BANK_LABEL}"
        )

        if st.button(
            hardest_label,
            key="select_hardest_bank",
            use_container_width=True,
            disabled=(is_hardest_selected or hardest_total_preview == 0),
        ):

            try:
                start_hardest_queue_in_session(
                    st.session_state.participant_id
                )
                st.rerun()

            except Exception as error:
                st.error("Could not build the hardest-questions deck.")
                st.exception(error)

        if hardest_total_preview == 0:

            st.caption(
                "Not enough answered questions yet to rank "
                "difficulty."
            )

        else:

            if is_hardest_selected:
                hardest_done = min(
                    st.session_state.base_completed
                    + st.session_state.current_index,
                    hardest_total_preview
                )
            else:
                hardest_done = 0

            st.progress(
                min(hardest_done / hardest_total_preview, 1.0)
            )

            st.caption(
                f"{hardest_done} / {hardest_total_preview} this "
                "session · lowest global accuracy"
            )

        st.divider()

        st.write("**Question Banks**")

        for bank in ALL_BANKS:

            bank_id = bank["bank_id"]

            if bank["error"]:
                st.caption(f"⚠️ {bank['display_name']} — failed to load")
                continue

            is_selected = st.session_state.selected_bank == bank_id

            label = f"{'▶ ' if is_selected else ''}{bank['display_name']}"

            if st.button(
                label,
                key=f"select_bank_{bank_id}",
                use_container_width=True,
                disabled=is_selected,
            ):

                try:
                    start_queue_in_session(
                        st.session_state.participant_id,
                        bank_id
                    )
                    st.rerun()

                except Exception as error:
                    st.error(
                        f"Could not load progress for "
                        f"{bank['display_name']}."
                    )
                    st.exception(error)

            try:
                cycle, done, total = get_bank_progress_summary(
                    st.session_state.participant_id,
                    bank_id
                )
            except Exception:
                cycle, done, total = 1, 0, len(bank["questions"])

            fraction = (done / total) if total else 0.0

            st.progress(min(fraction, 1.0))

            cycle_note = f" · cycle {cycle}" if cycle > 1 else ""

            st.caption(f"{done} / {total} completed{cycle_note}")

        st.divider()

        if st.session_state.selected_bank == HARDEST_BANK_ID:

            st.caption(
                "The hardest-questions deck always reflects "
                "live data, so there's no cycle to reset here — "
                "just reopen it above for a fresh set."
            )

        elif st.button(
            "Reset Progress for This Bank",
            use_container_width=True,
            help=(
                "Abandon the current cycle for the selected "
                "bank and start a brand-new randomized pass "
                "through it."
            )
        ):

            reset_participant_progress(
                st.session_state.participant_id,
                st.session_state.selected_bank
            )

            start_queue_in_session(
                st.session_state.participant_id,
                st.session_state.selected_bank
            )

            st.rerun()

    else:
        st.caption(f"{len(VALID_BANKS)} question bank(s) available.")


# ============================================================
# TITLE
# ============================================================

st.title("🧬 NBME Study Quiz")

st.caption(
    "Choose an answer. Incorrect answers can be "
    "retried until the correct answer is selected. "
    "Every question in the bank is asked once per cycle, "
    "in a random order, before anything repeats."
)


# ============================================================
# LOGIN / START PAGE
# ============================================================

if not st.session_state.quiz_started:

    st.subheader("Log in to start")

    username = st.text_input(
        "Username",
        placeholder="Enter your username"
    )

    password = st.text_input(
        "Password",
        type="password",
        placeholder="Enter your password"
    )

    if st.button(
        "Log In & Start Quiz",
        type="primary",
        use_container_width=True
    ):

        username = username.strip()

        users = st.secrets["users"]

        if username not in users:
            st.error("Invalid username or password.")
            st.stop()

        expected_password = users[username]

        if password != expected_password:
            st.error("Invalid username or password.")
            st.stop()

        participant_id = create_participant_id(
            username,
            st.secrets["PARTICIPANT_SALT"]
        )

        try:
            save_participant(participant_id, username)
        except Exception as error:
            st.error("Could not connect to the database.")
            st.exception(error)
            st.stop()

        st.session_state.quiz_started = True
        st.session_state.participant_id = participant_id
        st.session_state.display_name = username

        default_bank_id = VALID_BANKS[0]["bank_id"]

        try:
            start_queue_in_session(participant_id, default_bank_id)
        except Exception as error:
            st.error("Could not load your progress from the database.")
            st.exception(error)
            st.stop()

        st.rerun()

    st.divider()

    st.write("Please use the username and password provided to you.")

    st.stop()


# ============================================================
# ADMIN PANEL
# ============================================================

st.markdown(
    """
    <details>
    <summary><strong>Administrator</strong></summary>
    """,
    unsafe_allow_html=True
)

admin_password = st.text_input(
    "Administrator password",
    type="password",
    key="admin_password_input"
)

if admin_password:

    expected_password = st.secrets["ADMIN_PASSWORD"]

    if admin_password == expected_password:

        st.session_state.is_admin = True

        st.success("Administrator access granted.")

        try:

            all_answers = get_all_answers()

            if all_answers:

                admin_df = pd.DataFrame(all_answers)

                st.write(
                    f"Total recorded answer attempts: {len(admin_df):,}"
                )

                st.dataframe(admin_df, use_container_width=True)

                csv_data = admin_df.to_csv(index=False).encode("utf-8")

                st.download_button(
                    label="Download CSV",
                    data=csv_data,
                    file_name="nbme_quiz_results.csv",
                    mime="text/csv"
                )

            else:
                st.info("No answer records yet.")

        except Exception as error:
            st.error("Could not load administrator data.")
            st.exception(error)

        st.divider()

        st.write("**🚩 Flagged Questions**")

        try:

            flags = get_flagged_questions()

        except Exception as error:

            flags = []
            st.error("Could not load flagged questions.")
            st.exception(error)

        if flags:

            flags_by_bank_question = {}

            for flag in flags:

                key = (flag.get("bank_id"), flag["question_number"])

                flags_by_bank_question.setdefault(key, []).append(flag)

            sorted_keys = sorted(
                flags_by_bank_question.keys(),
                key=lambda item: (
                    BANKS_BY_ID.get(item[0], {}).get(
                        "display_name", item[0] or ""
                    ),
                    item[1]
                )
            )

            for flag_bank_id, q_number in sorted_keys:

                entries = flags_by_bank_question[(flag_bank_id, q_number)]

                bank = BANKS_BY_ID.get(flag_bank_id)
                bank_label = (
                    bank["display_name"] if bank
                    else (flag_bank_id or "Unknown bank")
                )

                col_info, col_trash = st.columns([4, 1])

                with col_info:
                    st.write(
                        f"**{bank_label}** — question **{q_number}** "
                        f"— flagged {len(entries)} time(s)"
                    )

                with col_trash:

                    if bank is not None and st.button(
                        "🗑️ Remove",
                        key=f"admin_trash_{flag_bank_id}_{q_number}"
                    ):

                        try:

                            remove_question_everywhere(
                                flag_bank_id,
                                q_number
                            )

                            if st.session_state.selected_bank == flag_bank_id:
                                drop_question_from_session(q_number)

                            st.success(
                                f"Question {q_number} removed from "
                                f"{bank_label}."
                            )

                            st.rerun()

                        except Exception as error:

                            st.error(
                                "Could not remove the question."
                            )

                            st.exception(error)

        else:
            st.caption("No questions have been flagged.")

    else:
        st.error("Incorrect administrator password.")

st.markdown("</details>", unsafe_allow_html=True)


# ============================================================
# CURRENT QUIZ
# ============================================================

selected_bank_id = st.session_state.selected_bank
is_hardest_mode = selected_bank_id == HARDEST_BANK_ID

if is_hardest_mode:
    current_bank_display_name = HARDEST_BANK_LABEL
else:
    current_bank_display_name = BANKS_BY_ID[selected_bank_id]["display_name"]

questions = st.session_state.questions

current_index = st.session_state.current_index

total_questions_all = st.session_state.selected_bank_total

st.caption(f"**Current bank:** {current_bank_display_name}")

if total_questions_all == 0:
    st.warning("There are no questions left in this question bank.")
    st.stop()


# ============================================================
# FINISHED
# ============================================================

if st.session_state.quiz_complete:

    st.success("🎉 Quiz complete!")

    if is_hardest_mode:
        st.write(
            f"You've gone through all **{total_questions_all}** "
            "of the current hardest questions. Nice work drilling "
            "your weak spots!"
        )
    else:
        st.write(
            f"You've completed cycle **{st.session_state.cycle}** of "
            f"**{current_bank_display_name}** — all "
            f"**{total_questions_all}** questions."
        )

    st.divider()

    st.info(
        "Your historical statistics are shown for each "
        "individual question while you are answering it."
    )

    restart_label = (
        "Pull a fresh hardest-questions deck"
        if is_hardest_mode
        else "Start another randomized quiz"
    )

    if st.button(
        restart_label,
        type="primary",
        use_container_width=True
    ):

        if is_hardest_mode:
            start_hardest_queue_in_session(
                st.session_state.participant_id
            )
        else:
            start_queue_in_session(
                st.session_state.participant_id,
                selected_bank_id
            )

        st.rerun()

    st.stop()


# ============================================================
# CURRENT QUESTION
# ============================================================

if current_index >= len(questions):
    # Safety net: the queue emptied out unexpectedly
    # (e.g. the last remaining question was just removed).
    st.session_state.quiz_complete = True
    st.rerun()

question = questions[current_index]
question_number = question["number"]

# When drilling the virtual "Hardest Questions" deck, every
# real action (saving an answer, flagging, removing) needs to
# route back to the bank the question actually came from —
# not to the virtual "hardest_questions" id, which has no real
# document or progress row behind it.
effective_bank_id = question.get("source_bank_id", selected_bank_id)


# ============================================================
# PROGRESS (consistent with overall cycle progress, not
# just position within this session's remaining queue)
# ============================================================

overall_position = min(
    st.session_state.base_completed + current_index + 1,
    total_questions_all
)

progress = overall_position / total_questions_all

st.progress(min(progress, 1.0))

st.write(f"### Question {overall_position} of {total_questions_all}")

if is_hardest_mode:

    source_bank_label = BANKS_BY_ID.get(
        effective_bank_id, {}
    ).get("display_name", effective_bank_id)

    st.caption(
        f"From **{source_bank_label}** · "
        f"Document question number: {question_number} · "
        f"Global accuracy: {question.get('source_accuracy', 0) * 100:.0f}% "
        f"({question.get('source_attempts', 0)} attempts)"
    )

else:

    st.caption(
        f"Cycle {st.session_state.cycle} · "
        f"Document question number: {question_number}"
    )


# ============================================================
# HISTORICAL QUESTION STATISTICS
# ============================================================

try:

    historical_stats = get_question_stats(
        st.session_state.participant_id,
        effective_bank_id,
        question_number
    )

except Exception:

    historical_stats = {
        "attempts": 0,
        "correct": 0,
        "accuracy": 0,
        "answers": {"A": 0, "B": 0, "C": 0, "D": 0, "E": 0},
        "global_attempts": 0,
        "global_correct": 0,
        "global_accuracy": 0,
    }

historical_attempts = historical_stats["attempts"]
historical_accuracy = historical_stats["accuracy"]

global_attempts = historical_stats["global_attempts"]
global_accuracy = historical_stats["global_accuracy"]


if st.session_state.last_feedback:

    if historical_attempts > 0:

        distribution_parts = []

        for letter in ["A", "B", "C", "D", "E"]:

            count = historical_stats["answers"].get(letter, 0)

            percentage = count / historical_attempts * 100

            distribution_parts.append(f"{letter}: {percentage:.0f}%")

        if global_attempts > 0:

            st.caption(
                f"Your previous attempts: **{historical_attempts}** | "
                f"Your historical accuracy: **{historical_accuracy:.0f}%** | "
                f"Global historical accuracy: **{global_accuracy:.0f}%** | "
                f"Previously selected: {' '.join(distribution_parts)}"
            )

        else:

            st.caption(
                f"Your previous attempts: **{historical_attempts}** | "
                f"Your historical accuracy: **{historical_accuracy:.0f}%** | "
                f"Global historical accuracy: No previous attempts"
            )

    else:

        if global_attempts > 0:

            st.caption(
                "Your previous attempts: **none** | "
                f"Global historical accuracy: **{global_accuracy:.0f}%**"
            )

        else:

            st.caption(
                "Your previous attempts: none | "
                "Global historical accuracy: no attempts yet"
            )

# ============================================================
# REPORT / REMOVE ROW
# ============================================================

flag_col, trash_col = st.columns([3, 1])

with flag_col:

    already_flagged = (
        (effective_bank_id, question_number)
        in st.session_state.flagged_this_session
    )

    if already_flagged:

        st.caption("🚩 Reported — thank you")

    else:

        if st.button(
            "🚩 Report question as irrelevant",
            key=f"flag_{effective_bank_id}_{question_number}"
        ):

            try:

                save_flag(
                    st.session_state.participant_id,
                    st.session_state.display_name,
                    effective_bank_id,
                    question_number
                )

                st.session_state.flagged_this_session.add(
                    (effective_bank_id, question_number)
                )

                st.rerun()

            except Exception as error:

                st.error("Could not save your report.")
                st.exception(error)

with trash_col:

    if st.session_state.is_admin:

        if st.button(
            "🗑️ Remove question",
            key=f"trash_{effective_bank_id}_{question_number}"
        ):

            try:

                remove_question_everywhere(
                    effective_bank_id,
                    question_number
                )

                drop_question_from_session(question_number)

                st.success(f"Question {question_number} removed.")

                st.rerun()

            except Exception as error:

                st.error("Could not remove the question.")
                st.exception(error)


# ============================================================
# QUESTION
# ============================================================

st.markdown(f"### {question['question']}")


# ============================================================
# ANSWER BUTTONS
# ============================================================

for letter in ["A", "B", "C", "D", "E"]:

    if letter not in question["options"]:
        continue

    option_text = question["options"][letter]
    button_text = f"{letter}. {option_text}"

    if st.session_state.current_correct:

        if letter == question["correct_letter"]:
            st.success(button_text)
        else:
            st.button(
                button_text,
                disabled=True,
                use_container_width=True,
                key=f"disabled_{selected_bank_id}_{current_index}_{letter}"
            )

        continue

    if st.button(
        button_text,
        use_container_width=True,
        key=f"question_{selected_bank_id}_{current_index}_{letter}"
    ):

        try:

            if is_hardest_mode:
                # The hardest-questions deck has no cycle of
                # its own — save against whatever cycle the
                # participant is actually on in the real bank
                # this question came from.
                effective_cycle = get_participant_cycle(
                    st.session_state.participant_id,
                    effective_bank_id
                )
            else:
                effective_cycle = st.session_state.cycle

            is_correct = save_answer(
                st.session_state.participant_id,
                st.session_state.display_name,
                effective_bank_id,
                question,
                letter,
                effective_cycle
            )

        except Exception as error:

            st.error("Your answer could not be saved.")
            st.exception(error)
            st.stop()

        if is_correct:

            st.session_state.current_correct = True
            st.session_state.last_feedback = "✓ Correct!"
            st.session_state.last_feedback_type = "correct"
            st.session_state.show_explanation = True
            st.session_state.flash_type = "correct"
            st.session_state.wrong_answer_streak = 0
            st.session_state.meltdown_popup = False

        else:

            st.session_state.last_feedback = "✗ Incorrect — try again."
            st.session_state.last_feedback_type = "incorrect"
            st.session_state.flash_type = "incorrect"
            st.session_state.wrong_answer_streak += 1

            if st.session_state.wrong_answer_streak >= 2:
                st.session_state.meltdown_popup = True

        st.rerun()


# ============================================================
# FEEDBACK
# ============================================================

if st.session_state.last_feedback:

    if st.session_state.last_feedback_type == "correct":
        st.success(st.session_state.last_feedback)
    elif st.session_state.last_feedback_type == "incorrect":
        st.error(st.session_state.last_feedback)


# ============================================================
# EXPLANATION
# ============================================================

if st.session_state.show_explanation:

    st.divider()

    st.subheader("Explanation")

    st.info(question["explanation"])

    st.write(
        f"**Correct answer: "
        f"{question['correct_letter']}. "
        f"{question['correct_answer']}**"
    )

    st.divider()

    is_last_in_queue = current_index == len(questions) - 1

    next_button_text = (
        "Finish Quiz" if is_last_in_queue else "Next Question →"
    )

    if st.button(
        next_button_text,
        type="primary",
        use_container_width=True
    ):

        if is_last_in_queue:
            st.session_state.quiz_complete = True
        else:
            st.session_state.current_index += 1

        st.session_state.current_correct = False
        st.session_state.last_feedback = ""
        st.session_state.last_feedback_type = ""
        st.session_state.show_explanation = False
        st.session_state.flash_type = None
        st.session_state.wrong_answer_streak = 0
        st.session_state.meltdown_popup = False

        st.rerun()
